import pandas as pd
import re
from langchain_community.chat_models import ChatOllama
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from collections import OrderedDict
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from flask import Flask, render_template, request, send_file

app = Flask(__name__)

#LLM model selected is wizardlm2
model_local = ChatOllama(model="wizardlm2")

#prompt is splitted into 2 proper generation of content

template = """
You are a Generative AI model resposible for generating Automotive Cybersecurity report on item definition for the item {question}.Generate maximum possible content for each of the subheadings given below, follow strict format no additional words or symbols in headers given below and all headers below should be included:

	Purpose:
		<Mention the purpose of the item definition document>
        
	Scope:
		<Define what is in scope and not in scope of item definition>
        
	Definitions, Acronyms and Abbreviations:
		<Generate as below format in table>
		Acronym 	Description

	Overview of Item:
		<Mention the details of the item under development, briefly describe the development strategies like newly developed, reused etc.>
               
	Components of the item:
		<Identify the components of the item and create the table in the below format>
		Item/Component	New development/Reuse/Reuse with modification	Remarks  
        
	Operational environment of the item:
		<Provide the description of the operational environment of the item in which the item will be operational.>
		<E.g. The environment/platform in which users run application software, user interface used by the program or item, etc.>     
    
	Item Boundary:
		<Identify the item boundary indicating all the external interfaces of the item. These interfaces can include interfaces internal and/or external to the vehicle.>
            
	Interaction with other Items/Component:
		<List all the other items / component with which the item under consideration will interact in a table in the format which has coulumns as Sl No,Items/ Component and Purpose>
			
	Internal Interfaces:
		<Elaborate the all internal components of the item and its interactions.>
            
	Assets of the Item:
		<Provide detailed description of the assets in item, identified to perform Threat analysis and Risk assessment in the table format below>
		 Sl No			Assets Identified			Functionality      
        
"""

template2="""
You are a Generative AI model resposible for generating Automotive Cybersecurity report on item definition for the item {question}.Generate maximum possible content for each of the subheadings given below, follow strict format no additional words or symbols in headers given below:

    Cybersecurity Requirements:
		<Provide description of the cybersecurity requirements for the item under consideration>
		
    Assumptions:
		<Assumptions about the item and the operational environment of the item shall be identified. Include assumptions on physical aspects and connectivity aspects.>
		<For Example: Assumptions on physical aspects can include the item will be placed in an anti-tamper enclosure.
		Assumptions on connectivity aspects can include every PKI certificate authority that the item relies on are appropriately managed.>

	Constraints and compliance:
		<Document the constraints and other standard compliance requirements. Constraints can include functional constraints, technical constraints etc. and compliances can include adherence to statutory and regulatory requirements and/or any other standards etc.>

	Known Vulnerabilities:
		<Provide details of the already known cybersecurity requirements, vulnerabilities>
"""

# form template with ChatPromptTemplate to give it as input to the model
after_prompt = ChatPromptTemplate.from_template(template)

# when below function is invoked the model executes to get output
after_chain = (
    {"question": RunnablePassthrough()}
    | after_prompt
    | model_local
    | StrOutputParser()
)

after_prompt2 = ChatPromptTemplate.from_template(template2)
after_chain2 = (
    {"question": RunnablePassthrough()}
    | after_prompt2
    | model_local
    | StrOutputParser()
)


def split_content_into_key_value(content):
    # Define regex patterns for headers
    header_pattern = re.compile(r'^(Purpose|Scope|Definitions, Acronyms and Abbreviations|Definitions, Acronyms, and Abbreviations|Overview of Item|Components of the item|Operational environment of the item|Item Boundary|Interaction with other Items/Component|Internal Interfaces|Assets of the Item|Cybersecurity Requirements|Assumptions|Constraints and Compliance|Known Vulnerabilities|Conclusion):$', re.MULTILINE)

    # Remove ** and # characters
    content = content.replace('*', '').replace('#', '')
    content = "\n".join(line.strip() for line in content.splitlines())

    # Split content based on headers
    headers = header_pattern.findall(content)
    splits = header_pattern.split(content)

    # Create an ordered dictionary to maintain the order of headers
    content_dict = OrderedDict()

    # Process the splits to populate the dictionary
    for i in range(1, len(splits), 2):
        header = splits[i].strip()
        value = splits[i + 1].strip()
        content_dict[header] = value

    return content_dict


def extract_table_and_content(value):
  table_data = ""
  remaining_content = ""

  for line in value.splitlines():
    if line.startswith('|'):
      table_data += line + "\n"  # Append table line with newline for formatting
    else:
      remaining_content += line + "\n"  # Append remaining line with newline

  return table_data


def set_cell_border(cell, **kwargs):
    tc_pr = cell._element.get_or_add_tcPr()
    for border_name, border_attrs in kwargs.items():
        border = OxmlElement(f'w:{border_name}')
        for attr_name, attr_val in border_attrs.items():
            border.set(qn(f'w:{attr_name}'), str(attr_val))
        tc_pr.append(border)
        

def add_table_to_doc(doc, table_text):
    lines = table_text.strip().split('\n')
    headers = lines[0].split('|')[1:-1]  # Extract headers
    data_lines = lines[2:]  # Extract data lines

    # Create a table in the document
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells

    # Fill in the headers
    for i, header in enumerate(headers):
        hdr_cells[i].text = header.strip()

    # Fill in the data rows
    for line in data_lines:
        data = line.split('|')[1:-1]
        row_cells = table.add_row().cells

        for i, item in enumerate(data):
            row_cells[i].text = item.strip()

        # Set border for each cell
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )

    return doc

def generate_document(item):
   print(f"Please wait for 10-15 mins, your artifact for {item} is generating.........")
   content1=after_chain.invoke(item)
   content2=after_chain2.invoke(item)
   content = content1+'\n\n'+content2
   print(content)

   document = Document()
   document.add_picture('logo.jpg', width=Inches(1))
   # Add title
   document.add_heading(f'{item} Automotive CyberSecurity Item Definition.', level=2)
   # Add metadata table
   menuTable = document.add_table(rows=7, cols=2)
   menuTable.style = 'Table Grid'
   hdr_Cells = menuTable.columns[0].cells
   hdr_Cells[0].text = 'Template Ref No:'
   hdr_Cells[1].text = 'Template Version:'
   hdr_Cells[2].text = 'Document No:'
   hdr_Cells[3].text = 'Document Ver:'
   hdr_Cells[4].text = 'Date:'
   hdr_Cells[5].text = 'Prepared by:'
   hdr_Cells[6].text = 'Approved by:'
   col_Cells = menuTable.columns[1].cells
   col_Cells[0].text = 'TDACS/QSTE/RE/01'
   col_Cells[1].text = '1.0'
   col_Cells[2].text = ''
   col_Cells[3].text = ''
   col_Cells[4].text = '9-mar-2021'
   col_Cells[5].text = 'AF'
   col_Cells[6].text = 'BN'
   # Add signature
   signature = document.add_paragraph("\nCompany Internal").alignment = WD_ALIGN_PARAGRAPH.RIGHT
   document.add_paragraph("Tata Elxsi").alignment = WD_ALIGN_PARAGRAPH.RIGHT
   # Add footer
   document.add_paragraph('\n\n\n\nCreated by Tata Elxsi., Quality Group\nCopyright 2021 Tata Elxsi.\nAll rights reserved.')\
       .alignment = WD_ALIGN_PARAGRAPH.CENTER
   document.add_paragraph('\nThis document contains information that is proprietary to Tata Elxsi. No part of this document may be reproduced\n or used in whole or part in any form or by any means - graphic, electronic or mechanical without the\n written permission of Tata Elxsi\n')\
       .alignment = WD_ALIGN_PARAGRAPH.CENTER
   # Add page break before data section
   document.add_page_break()
   # Load data from Excel
   data = pd.read_excel('sample_data.xlsx')
   # Add data table
   document.add_heading("Riview History \n\n")
   menuTable2 = document.add_table(rows=1, cols=5)
   menuTable2.style = 'Table Grid'
   hdr_Cells2 = menuTable2.rows[0].cells
   hdr_Cells2[0].text = 'Rev.No'
   hdr_Cells2[1].text = 'Date'
   hdr_Cells2[2].text = 'Author'
   hdr_Cells2[3].text = 'Reviewed/Approved by'
   hdr_Cells2[4].text = 'Description'
   # Populate data table
   for i in range(len(data)):
       cells = menuTable2.add_row().cells
       cells[0].text = str(data.iloc[i]['Rev.No'])
       cells[1].text = str(data.iloc[i]['Date'])
       cells[2].text = str(data.iloc[i]['Author'])
       cells[3].text = str(data.iloc[i]['Reviewed/Approved by'])
       cells[4].text = str(data.iloc[i]['Description'])
   document.add_page_break()
   
   
   content_dict = split_content_into_key_value(content)
   print("After split\n\n")
   for key, value in content_dict.items():
       print(f"{key}: \n\n {value}\n")
   
   for key, value in content_dict.items():
       print(f"{key},")
   
   for key, value in content_dict.items():
       document.add_heading(key, level=1)
       if '|' in value:  # Check if the value contains a table
	   
           table_data = extract_table_and_content(value)
           print(table_data)
           add_table_to_doc(document,table_data)
       else:
           document.add_paragraph(value)
   
   
   document.save('New.docx')


#Flask setup for running the route in browser
@app.route('/', methods=['GET'])
def index():
    return render_template('index.html', doc_generated=False)

@app.route('/generate', methods=['POST'])
def generate():
    item = request.form['item']
    generate_document(item)
    return render_template('index.html', doc_generated=True)

@app.route('/download')
def download_file():
    return send_file('New.docx', as_attachment=True, download_name='GeneratedDocument.docx')

if __name__ == '__main__':
    app.run(debug=True)